

from math import ceil
from sqlite3 import connect
import os
import time
import datetime
import os

from time import sleep,strptime,strftime

import openpyxl
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.styles import PatternFill,Font
from openpyxl.utils import get_column_letter
from openpyxl.styles.borders import Border, Side
from openpyxl.styles.alignment import Alignment

thin_border = Border(left=Side(style='thin'), 
                     right=Side(style='thin'), 
                     top=Side(style='thin'), 
                     bottom=Side(style='thin'))
center_alignment_obj = Alignment(horizontal='center', vertical='center')
right_alignment_obj = Alignment(horizontal='right', vertical='center')
wrap_text__alignment_obj = Alignment(wrap_text=True)

from re import sub

from docx import *
from docx.shared import Cm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_UNDERLINE
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT,WD_ALIGN_VERTICAL,WD_TABLE_ALIGNMENT

PACKET_SIZE =50


def set_row_height(row, height):
    #print(row.height)
    row.height = height
    for cell in row.cells:
        cell.height = height

def set_column_width(column, width):
    #print(column.width)
    column.width = width
    for cell in column.cells:
        cell.width = width

    
    

def coderintro():
    screen_time =3
    for i in range(0,screen_time):
                os.system('cls')
                print()
                print("********************************************************************************")
                print()
                print("developed by Keerthi Chandra C")
                print("Lecturer in ECE , GMRPW Karimnagar")
                print("for constructive criticism and suggestions reach me at keerthichand.c@gmail.com")
                print('''


disclaimer:
    *These are not actual TimeTables and NRs
    *Use this at your own discretion
    *for actual TimeTables please visit TS-SBTET Website
    *OUTPUTS from this program are not related to TS-SBTET
    *This program is developed only to help ease work but not to misguide or harm in 
        any possible ways and with no ill intent
    ''')
                print("******************************************************************************")
                if screen_time-i-1>0:
                        print(f'starting in {screen_time-i-2}s .......')
                sleep(1)
    input("press ENTER to start.....")
    os.system('cls')

def instructions_info():
        print('''
************************************************************************************

before staring ...
Make Sure that you
Paste your NRS ins respective excel files in the INPUTS folder
for more information read information.txt file

************************************************************************************''')
        input("press ENTER to continue......")
        os.system('cls')
        print()

def numtoWord(num):
    word=''
    numSet = ['ZERO','ONE','TWO','THREE','FOUR','FIVE','SIX','SEVEN','EIGHT','NINE']
    num=str(num)
    if int(num) <= 9:
        num = '0'+num
    for i in range(0,len(num)):
        word = word + numSet[int(num[i])]+' '

    
    return word
def simpleCommaPrint(a):
    if len(a)>0:
        s1=''
        for i in range(0,len(a)):
            j=i-1
            if j>=0:
                if a[i][0][2:5] ==  a[j][0][2:5]:
                    if a[i][0][0:2]==a[j][0][0:2]:
                        s1=s1+ ", "+a[i][0][-3:]
                        #print(","+a[i][0][-3:],end='')
                    else:
                        s1= s1 +'''
'''+ a[i][0]
                else :
                    s1 = s1+'''
'''+ a[i][0]
                    #print()
                    #print(a[i][0],end='')
            else:
                s1= s1 + '''
'''+ a[i][0]
                #print()
                #print(a[i][0],end='')
        #s1 = s1+ '''
         #         '''
        #print()
        #print(s1)
        return s1
    else:
        return '--NIL--'

def convListTossStrg(s):
    for S in s:
        return S[0]


def elDupes(a):
    return list(set(a))

def clean_str(to_str):
    #print(to_str)
    to_str = to_str.replace(" ","")
    to_str=to_str.replace("\n","")
    to_str=to_str.replace("\t","")
    to_str = to_str.replace("\r","")
    to_str = to_str.replace("\xa0","")
    return to_str


def clean_srting(to_str):
    to_str = to_str.replace(" ","")
    to_str=to_str.replace("\n","")
    to_str=to_str.replace("\t","")
    to_str = to_str.replace("\r","")
    to_str = to_str.replace("\xa0","")
    return to_str

def clean_subcode(to_str):
    to_str = sub(r'[^\w]', ' ', to_str)
    to_str = to_str.replace(" ","")
    to_str=to_str.replace("\n","")
    to_str=to_str.replace("\t","")
    to_str = to_str.replace("\r","")
    to_str = to_str.replace("\xa0","")
    return to_str


        


#os.chdir(os.path.dirname(__file__))
file_path = os.getcwd()
op_file_path=file_path+'\\OUTPUTS'

nr_table = 'nr_table'
time_table = 'time_table'
COLLEGE_INFO_TABLE = 'college_info_table'
template_excel = 'SeatingPlan_Template.xlsx'#'SeatingPlan_Empty_ReF_7.xlsx'#'SeatingPlan_Empty_ReF.xlsx'#'template_excel.xlsx'

nr_excels = ['C09_NR.xlsx','C14_NR.xlsx','C16_NR.xlsx','C16S_NR.xlsx','C18_NR.xlsx','C21_NR.xlsx']

conn = connect(file_path+'/srcs/exam_database.db')
curs = conn.cursor()

def get_from_DB(selection,tablename):
    a = curs.execute(f'''SELECT {selection} FROM {tablename};''')
    b =[]
    for (name,) in a:

        #name = elDupes(name)
        b.append(name)
    b= elDupes(b)
    b.sort()
    #print(b)
    return (b)







now = datetime.datetime.now()

tt_excels = ['TimeTable.xlsx']


def update_TT():
   
    curs.execute(f'drop table if exists {time_table}')
    conn.commit()
    curs.execute(f'''CREATE TABLE {time_table}(ID COUNTER PRIMARY KEY,scheme VARCHAR(10),session VARCHAR(10), date DATE,ExamTime VARCHAR(20), branch VARCHAR(10),sub_code_no VARCHAR(10),
										  sub_code VARCHAR(10),
										  sub_name VARCHAR(100),
										  year VARCHAR(20),
										  paper_code VARCHAR(10),
										  drawing VARCHAR(10));''')
    conn.commit()
    for file_name in tt_excels:
        print(f"Reading {file_name}.......",end='')
        wb = openpyxl.load_workbook(file_path+"\\INPUTS\\TimeTable\\"+file_name)
        sheet = wb.worksheets[0]
        a = []
        for rowNo in range(2,sheet.max_row):
            #print(rowNo)
            scheme = str(sheet.cell(row=rowNo,column=1).value)
            branch = str(sheet.cell(row=rowNo,column=2).value)
            paper_code = str(sheet.cell(row=rowNo,column=3).value)
            sub_code = sheet.cell(row=rowNo,column=4).value
            sub_name = str(sheet.cell(row=rowNo,column=5).value)
            date = str(sheet.cell(row=rowNo,column=6).value)
            ExamTime = str(sheet.cell(row=rowNo,column=7).value)
            session = str(sheet.cell(row=rowNo,column=8).value)
            ExamType = str(sheet.cell(row=rowNo,column=9).value)
            yr = str(sheet.cell(row=rowNo,column=10).value)
            sub_code_no = clean_subcode(sub_code[sub_code.find("-")+1:])
            sub_code_no =sub(r'[a-zA-Z]', '', sub_code_no)
            acdate = ""
            if 'DRAWING' in sub_name.upper():
                drawing = "YES"
            else:
                drawing = "NO"
            b = [date,acdate,session,scheme,branch,yr,sub_code,paper_code,sub_code_no,sub_name,drawing]
            #print(b)
            a.append(b)

        print('Done!')
        i=1
        print("Updatinging TimeTables in to DATABASAE.......",end = "")
        for [date,acdate,session,scheme,branch,yr,sub_code,paper_code,sub_code_no,sub_name,drawing] in a:
            
            print(i,date,session,scheme,branch,yr,sub_code,sub_code_no,sub_name,drawing)
            i+=1
            commd = f'''INSERT INTO {time_table}(scheme ,session,date,ExamTime,branch,sub_code_no,paper_code,sub_code,sub_name,year,drawing) VALUES ('{clean_str(scheme).upper()}','{clean_str(session).upper()}','{date}','{ExamTime}','{clean_str(branch).upper()}' ,'{clean_str(sub_code_no).upper()}','{paper_code}','{clean_str(sub_code).upper()}',"{clean_srting(sub_name)}",'{clean_str(yr).upper()}','{drawing.upper()}')'''
            #print(commd)
            curs.execute(commd)
    conn.commit()
    print("Done!")




def dateAndSession():
        
        now = datetime.datetime.now()
        print(now)
        if  int(now.hour) >= 9 and int(now.hour)<=12:
                Session = 'FN'
                Datesa=time.strftime("%Y-%m-%d")
        elif int(now.hour) >= 14 and int(now.hour)<=16:
                Session = 'AN'
                Datesa=time.strftime("%Y-%m-%d")
        else:
                Datesa = str(input('''Not Exam time !!!!! Enter Date Manually
                           Date in YYYY-MM-DD Format:?'''))
                print()
                Session = str(input('''Session??
                              AN for Afternoon
                              FN for Forenoon
                              Entrer : ''')).upper()
        return Datesa,Session






def noneType(ln):
    try:
        ln = str(ln).zfill(2)
    except:
        ln = '0'
    return ln





def create_nr():
    curs.execute(f'drop table if exists {nr_table}')
    conn.commit()
    curs.execute(f'''CREATE TABLE {nr_table}(ID COUNTER PRIMARY KEY,
                                          PinNo VARCHAR(20),
                                          student_name VARCHAR(50),
                                          Scheme VARCHAR(10),
                                          Branch VARCHAR(10),
                                          Year VARCHAR(10),
                                          SUB01 VARCHAR(10),
                                          SUB02 VARCHAR(10),
                                          SUB03 VARCHAR(10),
                                          SUB04 VARCHAR(10),
                                          SUB05 VARCHAR(10),
                                          SUB06 VARCHAR(10),
                                          SUB07 VARCHAR(10),
                                          SUB08 VARCHAR(10),
                                          SUB09 VARCHAR(10),
                                          SUB10 VARCHAR(10),
                                          SUB11 VARCHAR(10),
                                          SUB12 VARCHAR(10));''')
    conn.commit()

def create_TT():
    curs.execute(f'drop table if exists {time_table}')
    conn.commit()
    curs.execute(f'''CREATE TABLE {time_table}(ID COUNTER PRIMARY KEY,                              
                                          scheme VARCHAR(10),
                                          session VARCHAR(10),
                                          date DATE,
                                          branch VARCHAR(10),
                                          sub_code_no VARCHAR(10),
                                          sub_code VARCHAR(10),
                                          sub_name VARCHAR(100),
                                          year VARCHAR(20),
                                          drawing VARCHAR(10));''')
    conn.commit()
def create_col_info():
    curs.execute(f'drop table if exists {COLLEGE_INFO_TABLE}')
    conn.commit()
    curs.execute(f'''CREATE TABLE {COLLEGE_INFO_TABLE}(ID COUNTER PRIMARY KEY,                              
                                          collegeCode VARCHAR(20),
                                          collegeName VARCHAR(500));''')
    conn.commit()


def update_collInfo():
    create_col_info()

    code = str(input(f" Please Enter College Code: "))
    print()
    name = str(input(f" Please Enter College Name: "))

    curs.execute(f'''INSERT INTO {COLLEGE_INFO_TABLE}(collegeCode,collegeName) VALUES ('{code}','{name}')''')
    conn.commit()



def nr_DB():
    
    
    conn.commit()
    print("Updatinging NRs in to DATABASAE.......")
    for file_name in nr_excels:
        print(f"Reading {file_name}.......",end='')
        wb = openpyxl.load_workbook(file_path+"/INPUTS/NR_from_Principal_Login/"+file_name)
        sheet = wb.worksheets[0]
        cols= 18
        #print(sheet.max_row)
        for rowNo in range(1,sheet.max_row+2):
            b=[]
            #print(rowNo)
            pin = str(sheet.cell(row=rowNo,column=1).value).upper()
            sch = str(sheet.cell(row=rowNo,column=1+2).value).upper()
            if ('-' in pin) and ("C" in sch):
                for colNo in range(1,cols):
                    if colNo>5:
                            b.append(clean_srting(str(sheet.cell(row=rowNo,column=colNo).value).upper()))
                    else :
                        b.append(clean_srting(str(sheet.cell(row=rowNo,column=colNo).value).upper()))

                #print(b)
                #print(b[0],b[1],b[2],b[3],b[4],b[5],b[6],b[7],b[8],b[9],b[10],b[11],b[12],b[13],b[14],b[15])
                curs.execute('''INSERT INTO {}(PinNo ,student_name,Scheme,Year,Branch,SUB01,SUB02,SUB03,SUB04,SUB05,SUB06,SUB07,SUB08,SUB09,SUB10,SUB11,SUB12) VALUES ('{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}','{}')'''.format(nr_table,b[0],b[1],b[2],b[3],b[4],b[5],b[6],b[7],b[8],b[9],b[10],b[11],b[12],b[13],b[14],b[15],b[16]))
        print(" done!")
        conn.commit()
def deFill(slno,number_0f_scripts,number_0f_packets,date,session,scheme,dept,collegeName,collegeCode):
                pat = op_file_path+f'/DE-4 Sheets//{date}-{session}'
        
                if session == 'AN':
                
                    tim = ' 2:30 PM to 4:30 PM '
                
                elif session == 'FN':
                    tim = ' 9:30 AM to 11:30 AM '
                elif session == 'BN':
                    tim = ' 12:00 PM to 02:00 AM '
                
                          
                command = f'''SELECT year,sub_code,sub_name,sub_code_no FROM {time_table} WHERE date LIKE '{date}' 
                                                                                                    AND session LIKE '{session}'
                                                                                                    AND branch LIKE '{dept}'
                                                                                                    AND scheme LIKE '{scheme}';'''
                #print(command)
                curs.execute(command)
                a= curs.fetchall()
        
        
                if len(a) > 0:
                    a.sort()
                    a = elDupes(a)
                    
                    for (year,sub_code,sub_name,sub_code_no) in a:
                            command = f'''SELECT PinNo FROM {nr_table} WHERE scheme LIKE '{scheme}'
                                                                        AND branch LIKE '{dept}'
                                                                        AND year LIKE '{year}'
                                                                        AND (SUB01 LIKE '%{sub_code_no}%'
                                                                        OR SUB02 LIKE '%{sub_code_no}%'
                                                                        OR SUB03 LIKE '%{sub_code_no}%'
                                                                        OR SUB04 LIKE '%{sub_code_no}%'
                                                                        OR SUB05 LIKE '%{sub_code_no}%'
                                                                        OR SUB06 LIKE '%{sub_code_no}%'
                                                                        OR SUB07 LIKE '%{sub_code_no}%'
                                                                        OR SUB08 LIKE '%{sub_code_no}%'
                                                                        OR SUB09 LIKE '%{sub_code_no}%'
                                                                        OR SUB10 LIKE '%{sub_code_no}%'
                                                                        OR SUB11 LIKE '%{sub_code_no}%'
                                                                        OR SUB12 LIKE '%{sub_code_no}%')
                                                                        ORDER BY PinNo ASC;'''
                        
                            
                            curs.execute(command)
                            PinNumber = curs.fetchall()
                            print(PinNumber)
                            if len(PinNumber)> 0:
                                    PinNumber = elDupes(PinNumber)
                                    #print(pins)
                                    PinNumber.sort()
                                    
                                    #DE_document = Document(file_path+'/srcs/Template_Attendence.docx')
                                    DE_document = Document()

                                    DE_document = Document()
                                    section = DE_document.sections[0]
                                    section.page_height = Cm(29.7)
                                    section.page_width = Cm(21)
                                    sections = DE_document.sections

                                    for section in sections:
                                        margin =0.5
                                        margin1 = 2
                                        section.top_margin = Cm(margin)
                                        section.bottom_margin = Cm(margin)
                                        section.left_margin = Cm(margin1)
                                        section.right_margin = Cm(margin1)



                                    DE_document.add_heading('DE-4\nState Board of Technical Education & Training\nTELANGANA :: Hyderabad\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    p=DE_document.add_paragraph()
                                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    runner = p.add_run(f'ATTENDANCE SHEET')
                                    runner.font.underline = WD_UNDERLINE.SINGLE
                                    runner.bold = True
                                    runner.font.size = Pt(14)

                                    asnum= int(input(f" Number Of Absents in {scheme}{dept} {year} {sub_code} : "))
                                    absnum=[]
                                    PinNumberReg=[]
                                    if asnum == 0:
                                        absnum.append(('-NIL-',))
                        
                                    else:
                                        for i in range(0,asnum):
                                               abspinnum= str(input(f"  Pin Numbers Of Absents in {scheme}{dept} {year} {sub_code} : ")).upper()
                                               while True:
                                                   if ((abspinnum,)) in PinNumber: 
                                                       PinNumber.remove((abspinnum,))
                                                       absnum.append((abspinnum,))
                                                       break
                                                   else:
                                                       print("Wrong PinNumber!!! Enter Again: " )
                                                       abspinnum= str(input(f"  Pin Numbers Of Absents in {scheme}{dept} {year} {sub_code} : ")).upper()
                               



                                    #DE_document.add_heading(f'{slno+1}')
                                    table = DE_document.add_table(rows=12, cols=4)
                                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    table.autofit = False
                                    table.style = 'Light Shading'
                                    

                                    set_column_width(table.columns[0], Cm(1))
                                    set_column_width(table.columns[1], Cm(6.5))
                                    set_column_width(table.columns[2], Cm(.5))
                                    set_column_width(table.columns[3], Cm(8.5))

                                    

                                    hdr_cells = table.rows[0].cells
                                    hdr_cells[0].text = '1.'
                                    hdr_cells[1].text = 'Examination Centre'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = f'{collegeCode}, {collegeName}'

                                    hdr_cells = table.rows[1].cells
                                    hdr_cells[0].text = '2.'
                                    hdr_cells[1].text = 'Month & Year of Examination'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = f' {now.strftime("%B")}, {time.strftime("%Y")}'.upper()

                                    hdr_cells = table.rows[2].cells
                                    hdr_cells[0].text = '3.'
                                    hdr_cells[1].text = 'Day & Date of Examination'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = f'{time.strftime("%d/%m/%Y")}, {now.strftime("%A")}'


                                    hdr_cells = table.rows[3].cells
                                    hdr_cells[0].text = '4.  '
                                    hdr_cells[1].text = 'Time of Examination'
                                    hdr_cells[2].text = ':'                                 
                                    hdr_cells[3].text = f'{tim}'


                                    hdr_cells = table.rows[4].cells
                                    hdr_cells[0].text = '5.  '
                                    hdr_cells[1].text = 'Scheme and Name of the Examination'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = f'{scheme}- {year[0]} {year[1:]}'


                                    hdr_cells = table.rows[5].cells
                                    hdr_cells[0].text = '6.  '
                                    hdr_cells[1].text = 'Code No. & Name of the Subjects'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = f'{sub_code}, {sub_name}'


                                    hdr_cells = table.rows[6].cells
                                    hdr_cells[0].text = '7.  '
                                    hdr_cells[1].text = 'PIN of  the  Candidates Present'
                                    hdr_cells[2].text = ':'
                                    if len(PinNumber)<=PACKET_SIZE:
                                        hdr_cells[3].text = f"{noneType(simpleCommaPrint(PinNumber))}"
                                    if len(PinNumber)>PACKET_SIZE:
                                        for i in range(0,ceil(len(PinNumber)/PACKET_SIZE)):
                                            
                                                #print(f'Packet-{i+1}({len(PinNumber[i*50:(i+1)*50])}/{len(PinNumber)}):')
                                                
                                                if i != ceil(len(PinNumber)/PACKET_SIZE)-1:
                                                    text = hdr_cells[3].paragraphs[0].add_run(f'Packet-{i+1}({len(PinNumber[i*PACKET_SIZE:(i+1)*PACKET_SIZE])}/{len(PinNumber)}):')
                                                    text.bold= True
                                                    text.underline = True
                                                    text.font.size = Pt(15)
                                                    text = hdr_cells[3].paragraphs[0].add_run(f'{noneType(simpleCommaPrint(PinNumber[i*PACKET_SIZE:(i+1)*PACKET_SIZE]))}\n\n')
                                                else:
                                                    text = hdr_cells[3].paragraphs[0].add_run(f'Packet-{i+1}({len(PinNumber[i*PACKET_SIZE:(i+1)*PACKET_SIZE])}/{len(PinNumber)}):')
                                                    text.bold= True
                                                    text.underline = True
                                                    text.font.size = Pt(15)     
                                                    text = hdr_cells[3].paragraphs[0].add_run(f'{noneType(simpleCommaPrint(PinNumber[i*PACKET_SIZE:(i+1)*PACKET_SIZE]))}')
                                                    break
                                                             
                                            
                                        #hdr_cells[3].text = text

                                        #hdr_cells[3].text = f'''Packet-1({len(PinNumber[:50])}/{len(PinNumber)}): {noneType(simpleCommaPrint(PinNumber[:50]))}\n\nPacket-2({len(PinNumber[50:])}/{len(PinNumber)}): {noneType(simpleCommaPrint(PinNumber[50:]))}'''
                                        


                                    hdr_cells = table.rows[7].cells
                                    hdr_cells[0].text = '8.  '
                                    hdr_cells[1].text = 'PIN  of  the Candidates absent'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = f'{noneType(simpleCommaPrint(absnum))}'


                                    hdr_cells = table.rows[8].cells
                                    hdr_cells[0].text = '9.  '
                                    hdr_cells[1].text = 'No. of Answer Scripts(in Figure & Words)'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = f'{len(PinNumber)} ({numtoWord(len(PinNumber))} ONLY)'
                                    number_0f_scripts+= len(PinNumber)


                                    hdr_cells = table.rows[9].cells
                                    hdr_cells[0].text = '10.  '
                                    hdr_cells[1].text = 'PIN of Candidates for whom SG Booklets are issued'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = '--NIL--'


                                    hdr_cells = table.rows[10].cells
                                    hdr_cells[0].text = '11.  '
                                    hdr_cells[1].text = 'PIN  of  Candidates (if any) who Resorted to Malpractice'
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = '--NIL--'


                                    hdr_cells = table.rows[11].cells
                                    hdr_cells[0].text = '12.  '
                                    hdr_cells[1].text = 'Remarks if any '
                                    hdr_cells[2].text = ':'
                                    hdr_cells[3].text = '--NIL--'

                                    


                                 
                                    DE_document.add_paragraph('''\n\nSIGNATURE OF CHIEF SUPERINTENDENT ''').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    #DE_document.add_page_break()
                                    pat = op_file_path+f'/DE-4 Sheets/{date}-{session}'
                                    os.makedirs(pat, exist_ok=True)
                                    DE_document.save(pat+f'/{scheme}_{dept}_{year}.docx')

                                    doc2_cells = table2.add_row().cells
                                    #print(slno)
                                    doc2_cells[0].text = f'{slno+1}'
                                    doc2_cells[1].text = '{0}'.format('D'+dept+'E')
                                    doc2_cells[2].text = '{0}'.format(scheme)
                                    doc2_cells[3].text = '{0}'.format(sub_code)
                                    doc2_cells[4].text = f'{year[0]} {year[1:]}'
                                    doc2_cells[5].text = '{0}'.format(len(PinNumber))
                                    if len(PinNumber)== 0:
                                        doc2_cells[6].text = '{0}'.format(str(ceil(len(PinNumber)+1/PACKET_SIZE)))
                                    else:
                                        doc2_cells[6].text = '{0}'.format(str(ceil(len(PinNumber)/PACKET_SIZE)))
                                    doc2_cells[7].text = '{0}'.format('')

                                    do_cells = table1.rows[1].cells
                                    do_cells[0].text = 'Day & Date of Examination'
                                    do_cells[1].text = ':'
                                    do_cells[2].text = '{1}, {0}'.format(time.strftime("%d-%m-%Y"),now.strftime("%A"))

                                    do_cells = table1.rows[2].cells
                                    do_cells[0].text = 'Time of Examination'
                                    do_cells[1].text = ':'
                                    do_cells[2].text = '{0}'.format(tim)
                                    slno+=1
                                    if len(PinNumber)==0:
                                        number_0f_packets += 1
                                    else:    
                                        number_0f_packets += ceil(len(PinNumber)/PACKET_SIZE)

                return slno,number_0f_scripts,number_0f_packets,pat
                        
                
def summary_fill(collegeCode,collegeName,name_of_Exam):
    summary_document = Document()
    section = summary_document.sections[0]
    section.page_height = Cm(35.56)
    section.page_width = Cm(21.59)
    sections = summary_document.sections
    for section in sections:
                margin =1.27
                section.top_margin = Cm(margin)
                section.bottom_margin = Cm(margin)
                section.left_margin = Cm(margin)
                section.right_margin = Cm(margin)


    summary_document.add_heading(f'{collegeName}\n{name_of_Exam}\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p=summary_document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(f'Time Tables and NR Summary')
    runner.bold = True
    runner.font.size = Pt(95)
    summary_document.add_page_break()

    invisilation_list = openpyxl.Workbook()
    inv_sheet = invisilation_list.active
    inv_row = 4
    inv_col = 1
    inv_sheet.merge_cells(start_row=inv_row, start_column=inv_col, end_row=inv_row, end_column=inv_col+3)
    inv_sheet.cell(row=inv_row, column=inv_col).value =f'DATE'  
    inv_sheet.cell(row=inv_row, column=inv_col).font = Font(bold=True)
    inv_sheet.cell(row=inv_row, column=inv_col).border = thin_border
    inv_sheet.cell(row=inv_row, column=inv_col).alignment = right_alignment_obj
    inv_sheet.merge_cells(start_row=inv_row+1, start_column=inv_col, end_row=inv_row+1, end_column=inv_col+3)
    inv_sheet.cell(row=inv_row+1, column=inv_col).value =f'SESSION' 
    inv_sheet.cell(row=inv_row+1, column=inv_col).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+1, column=inv_col).alignment = right_alignment_obj
    inv_sheet.merge_cells(start_row=inv_row+2, start_column=inv_col, end_row=inv_row+2, end_column=inv_col+3)
    inv_sheet.cell(row=inv_row+2, column=inv_col).value =f'STRENGTH'    
    inv_sheet.cell(row=inv_row+2, column=inv_col).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+2, column=inv_col).alignment = right_alignment_obj
    inv_sheet.merge_cells(start_row=inv_row+3, start_column=inv_col, end_row=inv_row+3, end_column=inv_col+3)
    inv_sheet.cell(row=inv_row+3, column=inv_col).value =f'Aprox. No. of Invsilators'   
    inv_sheet.cell(row=inv_row+3, column=inv_col).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+3, column=inv_col).alignment = right_alignment_obj
    inv_sheet.merge_cells(start_row=inv_row+4, start_column=inv_col, end_row=inv_row+4, end_column=inv_col+3)
    inv_sheet.cell(row=inv_row+4, column=inv_col).value =f'Actual. No. of Invsilators'  
    inv_sheet.cell(row=inv_row+4, column=inv_col).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+4, column=inv_col).alignment = right_alignment_obj
    inv_sheet.cell(row=inv_row+5, column=inv_col).value =f'Sl.No' 
    inv_sheet.cell(row=inv_row+5, column=inv_col).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+5, column=inv_col).alignment = center_alignment_obj
    inv_sheet.cell(row=inv_row+5, column=inv_col+1).value =f'Name' 
    inv_sheet.cell(row=inv_row+5, column=inv_col+1).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+5, column=inv_col+1).alignment = center_alignment_obj
    inv_sheet.cell(row=inv_row+5, column=inv_col+2).value =f'Designation' 
    inv_sheet.cell(row=inv_row+5, column=inv_col+2).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+5, column=inv_col+2).alignment = center_alignment_obj

    inv_sheet.cell(row=inv_row+5, column=inv_col+3).value =f'Type of Duty' 
    inv_sheet.cell(row=inv_row+5, column=inv_col+3).font = Font(bold=True)
    inv_sheet.cell(row=inv_row+5, column=inv_col+3).alignment = center_alignment_obj
    inv_sheet.cell(row=inv_row+5, column=inv_col+3).alignment = wrap_text__alignment_obj

    inv_sheet.cell(row=inv_row+6, column=inv_col).value =f'1'
    #inv_sheet.cell(row=inv_row+6, column=inv_col).border = thin_border
    inv_sheet.cell(row=inv_row+6, column=inv_col).alignment = center_alignment_obj
    
    for i in range(1,22):
        inv_sheet.cell(row=inv_row+6+i, column=inv_col).value =f'=A{9+i}+1'
        inv_sheet.cell(row=inv_row+6+i, column=inv_col).border = thin_border
        inv_sheet.cell(row=inv_row+6+i, column=inv_col).alignment = center_alignment_obj
        

                
    #print(date_list)
    

    invisilation_rem = openpyxl.Workbook()

    inv_rem_sheet = invisilation_rem.active





    inv_rem_row = 4
    inv_rem_col = 1



    duties={ 'CS':225,
         'JCS':200,
         'IFS':160,
          'EDEP':120,
          'TABLE':120,
         'AS':120,
         'Invisilator':120,
         'Clerk':100,
         'Typist':80,
         'Attender': 82
            }

    itr =0
    f = ''
    for duty in duties: 
        itr+=1
        #ws[f'A{i}'].value= f'{duty}'
        if  itr != len(duties):
            f = f+f'IF(xxxfffttt11321="{duty}",{duties[duty]},'
        if itr == len(duties):
            f =f + f'IF(xxxfffttt11321="{duty}",{duties[duty]},""'+ itr*')'

    d = ''
    itr=0
    for key in duties.keys():
            itr+=1
            if itr!= len(duties):
                d = d+f'{key},'
            else:
                d = d +f'{key}'

    
    data_val = DataValidation(type="list",formula1=f'"{d}"',allow_blank=True) 
    inv_rem_sheet.add_data_validation(data_val)
    inv_sheet.add_data_validation(data_val)

    for itr in range(7,29):
        data_val.add(inv_rem_sheet[f"D{itr}"]) 
        data_val.add(inv_sheet[f"D{itr+3}"]) 
        
        inv_rem_sheet[f'E{itr}'].value= f"={f.replace('xxxfffttt11321',f'D{itr}')}"
        



    inv_rem_sheet.merge_cells(start_row=inv_rem_row, start_column=inv_rem_col, end_row=inv_rem_row, end_column=inv_rem_col+4)
    inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).value =f'DATE'  
    inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).font = Font(bold=True)
    inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).border = thin_border
    inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).alignment = right_alignment_obj
    inv_rem_sheet.merge_cells(start_row=inv_rem_row+1, start_column=inv_rem_col, end_row=inv_rem_row+1, end_column=inv_rem_col+4)
    inv_rem_sheet.cell(row=inv_rem_row+1, column=inv_rem_col).value =f'SESSION' 
    inv_rem_sheet.cell(row=inv_rem_row+1, column=inv_rem_col).font = Font(bold=True)
    inv_rem_sheet.cell(row=inv_rem_row+1, column=inv_rem_col).alignment = right_alignment_obj
    #inv_rem_sheet.merge_cells(start_row=inv_rem_row+2, start_column=inv_rem_col, end_row=inv_rem_row+2, end_column=inv_rem_col+4)
    # inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).value =f'STRENGTH'    
    # inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).font = Font(bold=True)
    # inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).alignment = right_alignment_obj
    #inv_rem_sheet.merge_cells(start_row=inv_rem_row+3, start_column=inv_rem_col, end_row=inv_rem_row+3, end_column=inv_rem_col+4)
    # inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).value =f'Aprox. No. of Invsilators'   
    # inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).font = Font(bold=True)
    # inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).alignment = right_alignment_obj
    # inv_rem_sheet.merge_cells(start_row=inv_rem_row+4, start_column=inv_rem_col, end_row=inv_rem_row+4, end_column=inv_rem_col+4)
    # inv_rem_sheet.cell(row=inv_rem_row+4, column=inv_rem_col).value =f'Actual. No. of Invsilators'  
    # inv_rem_sheet.cell(row=inv_rem_row+4, column=inv_rem_col).font = Font(bold=True)
    # inv_rem_sheet.cell(row=inv_rem_row+4, column=inv_rem_col).alignment = right_alignment_obj
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).value =f'Sl.No' 
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).font = Font(bold=True)
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).alignment = center_alignment_obj
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+1).value =f'Name' 
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+1).font = Font(bold=True)
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+1).alignment = center_alignment_obj
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+2).value =f'Designation' 
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+2).font = Font(bold=True)
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+2).alignment = center_alignment_obj


    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+3).value =f'Type of Duty ' 
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+3).font = Font(bold=True)
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+3).alignment = center_alignment_obj
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+3).alignment = wrap_text__alignment_obj

    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+4).value =f'Remuneration per Session' 
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+4).font = Font(bold=True)
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+4).alignment = center_alignment_obj
    inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col+4).alignment = wrap_text__alignment_obj

    inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).value =f'1'
    #inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).border = thin_border
    inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).alignment = center_alignment_obj


    for i in range(0,21):
        inv_rem_sheet.cell(row=inv_rem_row+4+i, column=inv_rem_col).value =f'=A{7+i}+1'
        inv_rem_sheet.cell(row=inv_rem_row+4+i, column=inv_rem_col).border = thin_border
        inv_rem_sheet.cell(row=inv_rem_row+4+i, column=inv_rem_col).alignment = center_alignment_obj


    coding_rem = openpyxl.Workbook()

    coding_rem_sheet = coding_rem.active





    coding_rem_row = 4
    coding_rem_col = 1



    duties={ 'Coding Officer': 0.25,
         'Attender':0.25,
         'Others': 0,
         
            }

    itr =0
    f = ''
    for duty in duties: 
        itr+=1
        #ws[f'A{i}'].value= f'{duty}'
        if  itr != len(duties):
            f = f+f'IF(xxxfffttt11321="{duty}",{duties[duty]},'
        if itr == len(duties):
            f =f + f'IF(xxxfffttt11321="{duty}",{duties[duty]},""'+ itr*')'

    d = ''
    itr=0
    for key in duties.keys():
            itr+=1
            if itr!= len(duties):
                d = d+f'{key},'
            else:
                d = d +f'{key}'

    
    data_val = DataValidation(type="list",formula1=f'"{d}"',allow_blank=True) 
    coding_rem_sheet.add_data_validation(data_val)
    inv_sheet.add_data_validation(data_val)

    for itr in range(7,31):
        data_val.add(coding_rem_sheet[f"D{itr}"]) 
        data_val.add(inv_sheet[f"D{itr+3}"]) 
        
        coding_rem_sheet[f'E{itr}'].value= f"={f.replace('xxxfffttt11321',f'D{itr}')}"
        



    coding_rem_sheet.merge_cells(start_row=coding_rem_row, start_column=coding_rem_col, end_row=coding_rem_row, end_column=coding_rem_col+4)
    coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).value =f'DATE'  
    coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).border = thin_border
    coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).alignment = right_alignment_obj
    coding_rem_sheet.merge_cells(start_row=coding_rem_row+1, start_column=coding_rem_col, end_row=coding_rem_row+1, end_column=coding_rem_col+4)
    coding_rem_sheet.cell(row=coding_rem_row+1, column=coding_rem_col).value =f'SESSION' 
    coding_rem_sheet.cell(row=coding_rem_row+1, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+1, column=coding_rem_col).alignment = right_alignment_obj
    coding_rem_sheet.merge_cells(start_row=coding_rem_row+2, start_column=coding_rem_col, end_row=coding_rem_row+2, end_column=coding_rem_col+4)
    coding_rem_sheet.cell(row=coding_rem_row+2, column=coding_rem_col).value =f'STRENGTH'    
    coding_rem_sheet.cell(row=coding_rem_row+2, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+2, column=coding_rem_col).alignment = right_alignment_obj
    coding_rem_sheet.merge_cells(start_row=coding_rem_row+3, start_column=coding_rem_col, end_row=coding_rem_row+3, end_column=coding_rem_col+4)
    coding_rem_sheet.cell(row=coding_rem_row+3, column=coding_rem_col).value =f'No. of OMRs Coded (total Present)'   
    coding_rem_sheet.cell(row=coding_rem_row+3, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+3, column=coding_rem_col).alignment = right_alignment_obj
    # coding_rem_sheet.merge_cells(start_row=coding_rem_row+4, start_column=coding_rem_col, end_row=coding_rem_row+4, end_column=coding_rem_col+4)
    # coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col).value =f'Actual. No. of Invsilators'  
    # coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col).font = Font(bold=True)
    # coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col).alignment = right_alignment_obj
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col).value =f'Sl.No' 
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col).alignment = center_alignment_obj
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+1).value =f'Name' 
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+1).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+1).alignment = center_alignment_obj
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+2).value =f'Designation' 
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+2).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+2).alignment = center_alignment_obj


    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+3).value =f'Type of Duty ' 
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+3).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+3).alignment = center_alignment_obj
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+3).alignment = wrap_text__alignment_obj

    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+4).value =f'Remuneration per OMR' 
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+4).font = Font(bold=True)
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+4).alignment = center_alignment_obj
    coding_rem_sheet.cell(row=coding_rem_row+4, column=coding_rem_col+4).alignment = wrap_text__alignment_obj

    coding_rem_sheet.cell(row=coding_rem_row+5, column=coding_rem_col).value =f'1'
    #coding_rem_sheet.cell(row=coding_rem_row+5, column=coding_rem_col).border = thin_border
    coding_rem_sheet.cell(row=coding_rem_row+5, column=coding_rem_col).alignment = center_alignment_obj


    for i in range(0,21):
        coding_rem_sheet.cell(row=coding_rem_row+6+i, column=coding_rem_col).value =f'=A{9+i}+1'
        coding_rem_sheet.cell(row=coding_rem_row+6+i, column=coding_rem_col).border = thin_border
        coding_rem_sheet.cell(row=coding_rem_row+6+i, column=coding_rem_col).alignment = center_alignment_obj
   
   

    inv_rem_col =6
    inv_col =5
    coding_rem_col = 6
    inv_s = inv_col

    print("Creating Day and Session Wise EXCELs.......")
    for date in date_list:

        
        if ((inv_col-6)>0 and ((inv_s-5)%6)==0):

            inv_sheet.merge_cells(start_row=inv_row, start_column=inv_col, end_row=inv_row+5, end_column=inv_col)
            
            inv_sheet.cell(row=inv_row, column=inv_col).value = 'Signature'
            inv_sheet.cell(row=inv_row, column=inv_col).font = Font(bold=True)
            inv_sheet.cell(row=inv_row, column=inv_col).border = thin_border
            inv_sheet.cell(row=inv_row, column=inv_col).alignment = center_alignment_obj
            inv_sheet.cell(row=inv_row, column=inv_col).alignment = wrap_text__alignment_obj
            inv_sheet.column_dimensions[get_column_letter(inv_col)].width = 12
            inv_col+=1
            #inv_s = inv_col-j
            inv_sheet.merge_cells(start_row=inv_row, start_column=inv_col, end_row=inv_row, end_column=inv_col+2)
            inv_sheet.cell(row=inv_row, column=inv_col).value = f'{strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}'
            inv_sheet.cell(row=inv_row, column=inv_col).font = Font(bold=True)
            inv_sheet.cell(row=inv_row, column=inv_col).border = thin_border
            inv_sheet.cell(row=inv_row, column=inv_col+1).border = thin_border
            inv_sheet.cell(row=inv_row, column=inv_col).alignment = center_alignment_obj

        else:
            inv_sheet.merge_cells(start_row=inv_row, start_column=inv_col, end_row=inv_row, end_column=inv_col+2)
            inv_sheet.cell(row=inv_row, column=inv_col).value = f'{strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}'
            inv_sheet.cell(row=inv_row, column=inv_col).font = Font(bold=True)
            inv_sheet.cell(row=inv_row, column=inv_col).border = thin_border
            inv_sheet.cell(row=inv_row, column=inv_col+1).border = thin_border
            inv_sheet.cell(row=inv_row, column=inv_col).alignment = center_alignment_obj


        inv_rem_sheet.merge_cells(start_row=inv_rem_row, start_column=inv_rem_col, end_row=inv_rem_row, end_column=inv_rem_col+2)
        inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).value = f'{strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}'
        inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).font = Font(bold=True)
        inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).border = thin_border
        inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col+1).border = thin_border
        inv_rem_sheet.cell(row=inv_rem_row, column=inv_rem_col).alignment = center_alignment_obj



        coding_rem_sheet.merge_cells(start_row=coding_rem_row, start_column=coding_rem_col, end_row=coding_rem_row, end_column=coding_rem_col+2)
        coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).value = f'{strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}'
        coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).font = Font(bold=True)
        coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).border = thin_border
        coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col+1).border = thin_border
        coding_rem_sheet.cell(row=coding_rem_row, column=coding_rem_col).alignment = center_alignment_obj
        
        for session in session_list:
            inv_sheet.cell(row=inv_row+1, column=inv_col).value = f'{session}'
            inv_sheet.cell(row=inv_row+1, column=inv_col).border = thin_border
            inv_sheet.cell(row=inv_row+1, column=inv_col).alignment = center_alignment_obj

            #inv_sheet.merge_cells(start_row=inv_row, start_column=inv_col, end_row=inv_row, end_column=inv_col+1)


            inv_rem_sheet.cell(row=inv_rem_row+1, column=inv_rem_col).value = f'{session}'
            inv_rem_sheet.cell(row=inv_rem_row+1, column=inv_rem_col).border = thin_border
            inv_rem_sheet.cell(row=inv_rem_row+1, column=inv_rem_col).alignment = center_alignment_obj



            coding_rem_sheet.cell(row=coding_rem_row+1, column=coding_rem_col).value = f'{session}'
            coding_rem_sheet.cell(row=coding_rem_row+1, column=coding_rem_col).border = thin_border
            coding_rem_sheet.cell(row=coding_rem_row+1, column=coding_rem_col).alignment = center_alignment_obj

            print(f"Filling PIN in EXCEL for DATE: {date} and Session: {session}............",end='')
            grand_total = 0
            ours_grand_total = 0
            others_grand_total = 0
            ours = 0
            others =0
            
            #print(f'        Date: {date}       Session : {session}')
            #document.add_heading(f'        Date: {date}       Session : {session}')
            summary_document.add_heading(f'{collegeName}').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            summary_document.add_heading(f'        Date: {strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}       Session : {session}')
            
            table = summary_document.add_table(rows=1, cols=5)
            topRow = ['Scheme','Year','Subject Code','Pin Numbers','Total']
            hdr_cells = table.rows[0].cells
            for i in range(0,len(hdr_cells)):
                                hdr_cells[i].text = topRow[i]
                                run = hdr_cells[i].paragraphs[0].runs[0]
                                run.font.bold = True
            table.autofit = False
            table.style = 'Light Shading'
            #document.save(file_path+'/op.docx')
            wb = openpyxl.load_workbook("srcs/"+template_excel)
            sheet = wb['MAINWORKING']
            sheet2 = wb['SEATINGPLAN']
            sheet3 = wb['DISPLAY']
            rowno =0
            colno =0
            for scheme in scheme_list:

                for dept in branch_list:
                    ours = 0
                    others =0
                    
                    command = f'''SELECT year,sub_code,paper_code,sub_name,sub_code FROM {time_table} WHERE date LIKE '{date}' 
                                                                                                    AND session LIKE '{session}'
                                                                                                    AND branch LIKE '{dept}'
                                                                                                    AND scheme LIKE '{scheme}';'''
                    #print(command)
                    curs.execute(command)
                    a= curs.fetchall()
                    #print(a)
                    if len(a) > 0:
                        for (year,sub_code,paper_code,sub_name,sub_code_no) in a:
                            rowno =1
                            colno +=1
                            #print(year,sub_code,sub_name,sub_code_no)
                             
                            command = f'''SELECT PinNo FROM {nr_table} WHERE scheme LIKE '{scheme}'
                                                                        AND branch LIKE '{dept}'
                                                                        AND year LIKE '{year}'
                                                                        AND (SUB01 LIKE '%{sub_code}%'
                                                                        OR SUB02 LIKE '%{sub_code}%'
                                                                        OR SUB03 LIKE '%{sub_code}%'
                                                                        OR SUB04 LIKE '%{sub_code}%'
                                                                        OR SUB05 LIKE '%{sub_code}%'
                                                                        OR SUB06 LIKE '%{sub_code}%'
                                                                        OR SUB07 LIKE '%{sub_code}%'
                                                                        OR SUB08 LIKE '%{sub_code}%'
                                                                        OR SUB09 LIKE '%{sub_code}%'
                                                                        OR SUB10 LIKE '%{sub_code}%'
                                                                        OR SUB11 LIKE '%{sub_code}%'
                                                                        OR SUB12 LIKE '%{sub_code}%')
                                                                        ORDER BY PinNo ASC;'''
                            
                                
                            curs.execute(command)
                            pins = curs.fetchall()
                            pins = elDupes(pins)
                            #print(pins)
                            pins.sort()
                            #print(pins)

                            row_cells = table.add_row().cells
                            row_cells[0].text = f'{scheme}'
                            row_cells[1].text = f'{year}'
                            row_cells[2].text = f'''{sub_code}(Paper Code: {paper_code})\n{sub_name}'''
                            row_cells[3].text = noneType(simpleCommaPrint(pins))
                            if len(pins) == 0:
                                row_cells[4].text = noneType(len(pins))
                            else:
                                ours = 0
                                others =0
                                for (pin,) in pins:
                                    if collegeCode in pin.split('-')[0]:
                                        ours+=1
                                    else:
                                        others+=1
                                if others == 0:
                                    row_cells[4].text = noneType(len(pins))
                                        
                                else:  
                                    row_cells[4].text = {noneType(len(pins))}
#ours:{ours}
#+
#Others:{others}'''

                            grand_total += len(pins)
                            ours_grand_total += ours
                            others_grand_total += others

                            sheet.cell(row=rowno, column=colno).value =f'{scheme} {sub("[^a-zA-Z]+", "", sub_code)}' 
                            sheet.cell(row=rowno, column=colno).font = Font(bold=True)
                            #sheet.cell(row=rowno, column=colno).fill = PatternFill(bgColor="FFC7CE", fill_type = "solid")

                            sheet.cell(row=rowno+1, column=colno).value =f'{sub_code_no}({len(pins)})'
                            sheet.cell(row=rowno+1, column=colno).font = Font(bold=True)
                            sheet.cell(row=rowno+2, column=colno).value =f'{sub_name}'
                            sheet.cell(row=rowno+2, column=colno).font = Font(bold=True)


                            sheet3.cell(row=rowno, column=2*colno-1).value =f'{scheme} {sub("[^a-zA-Z]+", "", sub_code)}' 
                            sheet3.cell(row=rowno, column=2*colno-1).font = Font(bold=True)
                            #sheet.cell(row=rowno, column=colno).fill = PatternFill(bgColor="FFC7CE", fill_type = "solid")

                            sheet3.cell(row=rowno+1, column=2*colno-1).value =f'{sub_code_no}({len(pins)})'
                            sheet3.cell(row=rowno+1, column=2*colno-1).font = Font(bold=True)
                            sheet3.cell(row=rowno+2, column=2*colno-1).value =f'{sub_name}'
                            sheet3.cell(row=rowno+2, column=2*colno-1).font = Font(bold=True)
                            #sheet.cell(row=rowno+1, column=colno).fill = PatternFill(bgColor="FFC7CE", fill_type = "solid")
                            if len(pins)>0:
                                rowno = 4
                                for (pin,) in pins:

                                    sheet.cell(row=rowno, column=colno).value =f'{pin}'
                                    
                                    sheet3.cell(row=rowno, column=2*colno-1).value =f'{pin}'
                                    #sheet3.cell(row=rowno, column=2*colno).value =f'''=INDEX(SHEET_LIST,MATCH(1,--(COUNTIF(INDIRECT("'"&SHEET_LIST&"'!$A$4:$A$30"),{get_column_letter(2*colno-1)}{rowno})>0),0))'''
                                    sheet3.cell(row=rowno, column=2*colno).value =f'''=LOOKUP(9.99999999999999E+307,1/COUNTIF(INDIRECT("'"&SHEET_LIST&"'!A2:O30"),${get_column_letter(2*colno-1)}{rowno}),SHEET_LIST)'''
                                    rowno +=1
                            else:
                                rowno = 4
                                sheet.cell(row=rowno, column=colno).value ='--NIL--'
                                sheet.cell(row=rowno, column=colno).font = Font(bold=True)

                                sheet3.cell(row=rowno, column=2*colno-1).value ='--NIL--'
                                sheet3.cell(row=rowno, column=2*colno-1).font = Font(bold=True)
                                colno -=1

            
            
            set_column_width(table.columns[0], Cm(1.8))
            set_column_width(table.columns[1], Cm(1.5))
            set_column_width(table.columns[2], Cm(4))
            set_column_width(table.columns[3], Cm(11))
            set_column_width(table.columns[4], Cm(1.4))
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(10)
            if others_grand_total == 0:
                summary_document.add_heading(f'                         Grand Total : {grand_total}')
            else:
                summary_document.add_heading(f'                         Grand Total : {grand_total}(ours:{ours_grand_total}+others:{others_grand_total})')
            paragraph = summary_document.add_paragraph()
            paragraph.text = '''           Developed by Keerthi Chandra C, Lecturer in ECE, GMRPW Karimnagar
                                        reach me at keerthichand.c@gmail.com'''



            summary_document.add_page_break()
            #sheet.column_dimensions.hidden = False

            for i in range(0,30):
                sheet.column_dimensions[get_column_letter(i+1)].width = 16
            sheet2.cell(row=4, column=3).value =f'{strftime("%d-%m-%Y",strptime(date,"%Y-%m-%d"))}'
            sheet2.cell(row=1, column=1).value =f'{name_of_Exam}'
            sheet2.cell(row=4, column=5).value =f'{session}'
            sheet2.cell(row=22, column=5).value =f'{grand_total}'
            os.makedirs(op_file_path, exist_ok=True)
            wb.save(op_file_path+f'\\dayWise_PIN_EXCEL\\{date}_{session}.xlsx')
            print("done!")
            inv_sheet.cell(row=inv_row+2, column=inv_col).value = f'{grand_total}'
            inv_sheet.cell(row=inv_row+2, column=inv_col).border = thin_border
            if ceil(grand_total/24)==1:
                no_invs_yo_day = 2
            else:
                no_invs_yo_day = ceil(grand_total/24)
            inv_sheet.cell(row=inv_row+3, column=inv_col).value = f'{no_invs_yo_day}'
            inv_sheet.cell(row=inv_row+3, column=inv_col).border = thin_border
            inv_col+=1
            inv_s +=1

            # inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).value = f'{grand_total}'
            # inv_rem_sheet.cell(row=inv_rem_row+2, column=inv_rem_col).border = thin_border

            coding_rem_sheet.cell(row=coding_rem_row+2, column=coding_rem_col).value = f'{grand_total}'
            coding_rem_sheet.cell(row=coding_rem_row+2, column=coding_rem_col).border = thin_border
            # inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).value = f'{ceil(grand_total/24)}'
            # inv_rem_sheet.cell(row=inv_rem_row+3, column=inv_rem_col).border = thin_border
            inv_rem_col+=1
            coding_rem_col+=1

        #inv_col+=1    


    print("Created all Excel Files Sucessfully!")
    print("Creating Total Exam summary.........",end='')
    os.makedirs(op_file_path, exist_ok=True)
    summary_document.save(op_file_path+'\\Total_Exam_summary.docx')

    inv_sheet.merge_cells(start_row=inv_row, start_column=inv_col, end_row=inv_row+5, end_column=inv_col)
            
    inv_sheet.cell(row=inv_row, column=inv_col).value = 'Signature'
    inv_sheet.cell(row=inv_row, column=inv_col).font = Font(bold=True)
    inv_sheet.cell(row=inv_row, column=inv_col).border = thin_border
    inv_sheet.cell(row=inv_row, column=inv_col).alignment = center_alignment_obj
    inv_sheet.cell(row=inv_row, column=inv_col).alignment = wrap_text__alignment_obj
    inv_sheet.column_dimensions[get_column_letter(inv_col)].width = 12 


    inv_rem_sheet.merge_cells(start_row=4, start_column=inv_rem_col, end_row=6, end_column=inv_rem_col)

    inv_rem_sheet.cell(row=4, column=inv_rem_col).value = 'Total Number of Sessions'
    inv_rem_sheet.cell(row=4, column=inv_rem_col).font = Font(bold=True)
    inv_rem_sheet.cell(row=4, column=inv_rem_col).border = thin_border
    inv_rem_sheet.cell(row=4, column=inv_rem_col).alignment = center_alignment_obj
    inv_rem_sheet.cell(row=4, column=inv_rem_col).alignment = wrap_text__alignment_obj
    inv_rem_sheet.column_dimensions[get_column_letter(inv_rem_col)].width = 12 

    inv_rem_sheet.merge_cells(start_row=4, start_column=inv_rem_col+1, end_row=6, end_column=inv_rem_col+1)

    inv_rem_sheet.cell(row=4, column=inv_rem_col+1).value = 'Total Remuneration'
    inv_rem_sheet.cell(row=4, column=inv_rem_col+1).font = Font(bold=True)
    inv_rem_sheet.cell(row=4, column=inv_rem_col+1).border = thin_border
    inv_rem_sheet.cell(row=4, column=inv_rem_col+1).alignment = center_alignment_obj
    inv_rem_sheet.cell(row=4, column=inv_rem_col+1).alignment = wrap_text__alignment_obj
    inv_rem_sheet.column_dimensions[get_column_letter(inv_rem_col+1)].width = 12 

    for i in range(0,22):
        inv_rem_sheet.cell(row=inv_rem_row+3+i, column=inv_rem_col).value =f'=SUM(F{inv_rem_row+3+i}:{get_column_letter(inv_rem_col-1)}{inv_rem_row+3+i})'
        inv_rem_sheet.cell(row=inv_rem_row+3+i, column=inv_rem_col).border = thin_border
        inv_rem_sheet.cell(row=inv_rem_row+3+i, column=inv_rem_col).alignment = center_alignment_obj
        
        inv_rem_sheet.cell(row=inv_rem_row+3+i, column=inv_rem_col+1).value =f'=E{inv_rem_row+3+i}*{get_column_letter(inv_rem_col)}{inv_rem_row+3+i}'
        
        inv_rem_sheet.cell(row=inv_rem_row+3+i, column=inv_rem_col+1).border = thin_border
        inv_rem_sheet.cell(row=inv_rem_row+3+i, column=inv_rem_col+1).alignment = center_alignment_obj



    inv_rem_sheet.merge_cells(start_row=29, start_column=1, end_row=29, end_column=5)
    inv_rem_sheet.cell(row=29, column=inv_rem_col).value =f'Total'  
    inv_rem_sheet.cell(row=29, column=inv_rem_col).font = Font(bold=True)
    inv_rem_sheet.cell(row=29, column=inv_rem_col).border = thin_border
    inv_rem_sheet.cell(row=29, column=inv_rem_col).alignment = right_alignment_obj

    inv_rem_sheet.cell(row=29, column=inv_rem_col+1).value =f'=SUM({get_column_letter(inv_rem_col+1)}{inv_rem_row+3}:{get_column_letter(inv_rem_col+1)}{28})'  
    inv_rem_sheet.cell(row=29, column=inv_rem_col+1).font = Font(bold=True)
    inv_rem_sheet.cell(row=29, column=inv_rem_col+1).border = thin_border
    inv_rem_sheet.cell(row=29, column=inv_rem_col+1).alignment = center_alignment_obj

    inv_rem_sheet.merge_cells(start_row=30, start_column=1, end_row=33, end_column=inv_rem_col+1)
    inv_rem_sheet.cell(row=30, column=1).value =f'PRINCIPAL'  
    inv_rem_sheet.cell(row=30, column=1).font = Font(bold=True)
    #inv_rem_sheet.cell(row=30, column=1).border = thin_border
    inv_rem_sheet.cell(row=30, column=1).alignment = right_alignment_obj



    for i in range(6,inv_rem_col):
        inv_rem_sheet.cell(row=29, column= i).value =f'=SUM({get_column_letter(i)}7:{get_column_letter(i)}28)'


    #inv_rem_sheet.insert_rows(1,1)
    inv_rem_sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=inv_rem_col+1)
    inv_rem_sheet.cell(row=1, column=1).value =f'{collegeName}'
    inv_rem_sheet.cell(row=1, column=1).font = Font(bold=True)
    inv_rem_sheet.cell(row=1, column=1).font = Font(size=16)
    #inv_rem_sheet.cell(row=1, column=1).border = thin_border
    inv_rem_sheet.cell(row=1, column=1).alignment = center_alignment_obj

    inv_rem_sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=inv_rem_col+1)
    inv_rem_sheet.cell(row=2, column=1).value =f'{name_of_Exam}'
    inv_rem_sheet.cell(row=2, column=1).font = Font(bold=True)
    inv_rem_sheet.cell(row=2, column=1).font = Font(size = 16)
    #inv_rem_sheet.cell(row=2, column=1).border = thin_border
    inv_rem_sheet.cell(row=2, column=1).alignment = center_alignment_obj


    inv_rem_sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=inv_rem_col+1)
    inv_rem_sheet.cell(row=3, column=1).value =f'Proforma showing remunerations for the spell of Examination'
    inv_rem_sheet.cell(row=3, column=1).font = Font(bold=True)
    inv_rem_sheet.cell(row=3, column=1).font = Font(size = 16)
    #inv_rem_sheet.cell(row=3, column=1).border = thin_border
    inv_rem_sheet.cell(row=3, column=1).alignment = center_alignment_obj

    inv_rem_sheet.merge_cells(start_row=6, start_column=6, end_row=6, end_column=inv_rem_col-1)


    inv_sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=inv_col)
    inv_sheet.cell(row=1, column=1).value =f'{collegeName}'
    inv_sheet.cell(row=1, column=1).font = Font(bold=True)
    inv_sheet.cell(row=1, column=1).font = Font(size=16)
    #inv_sheet.cell(row=1, column=1).border = thin_border
    inv_sheet.cell(row=1, column=1).alignment = center_alignment_obj

    inv_sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=inv_col)
    inv_sheet.cell(row=2, column=1).value =f'{name_of_Exam}'
    inv_sheet.cell(row=2, column=1).font = Font(bold=True)
    inv_sheet.cell(row=2, column=1).font = Font(size = 16)
    #inv_sheet.cell(row=2, column=1).border = thin_border
    inv_sheet.cell(row=2, column=1).alignment = center_alignment_obj


    inv_sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=inv_col)
    inv_sheet.cell(row=3, column=1).value =f'Proforma showing duties for the spell of Examination'
    inv_sheet.cell(row=3, column=1).font = Font(bold=True)
    inv_sheet.cell(row=3, column=1).font = Font(size = 16)
    #inv_sheet.cell(row=3, column=1).border = thin_border
    inv_sheet.cell(row=3, column=1).alignment = center_alignment_obj

    inv_sheet.merge_cells(start_row=32, start_column=1, end_row=35, end_column=inv_col)
    inv_sheet.cell(row=32, column=1).value =f'PRINCIPAL'  
    inv_sheet.cell(row=32, column=1).font = Font(bold=True)
    #inv_sheet.cell(row=32, column=1).border = thin_border
    inv_sheet.cell(row=32, column=1).alignment = right_alignment_obj


    coding_rem_sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=coding_rem_col)
    coding_rem_sheet.cell(row=1, column=1).value =f'{collegeName}'
    coding_rem_sheet.cell(row=1, column=1).font = Font(bold=True)
    coding_rem_sheet.cell(row=1, column=1).font = Font(size=16)
    #coding_rem_sheet.cell(row=1, column=1).border = thin_border
    coding_rem_sheet.cell(row=1, column=1).alignment = center_alignment_obj

    coding_rem_sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=coding_rem_col)
    coding_rem_sheet.cell(row=2, column=1).value =f'{name_of_Exam}'
    coding_rem_sheet.cell(row=2, column=1).font = Font(bold=True)
    coding_rem_sheet.cell(row=2, column=1).font = Font(size = 16)
    #coding_rem_sheet.cell(row=2, column=1).border = thin_border
    coding_rem_sheet.cell(row=2, column=1).alignment = center_alignment_obj


    coding_rem_sheet.merge_cells(start_row=3, start_column=1, end_row=3, end_column=coding_rem_col)
    coding_rem_sheet.cell(row=3, column=1).value =f'Proforma showing coding remunerations for the spell of Examination'
    coding_rem_sheet.cell(row=3, column=1).font = Font(bold=True)
    coding_rem_sheet.cell(row=3, column=1).font = Font(size = 16)
    #coding_rem_sheet.cell(row=3, column=1).border = thin_border
    coding_rem_sheet.cell(row=3, column=1).alignment = center_alignment_obj



    coding_rem_sheet.merge_cells(start_row=4, start_column=coding_rem_col, end_row=6, end_column=coding_rem_col)
    coding_rem_sheet.cell(row=4, column=coding_rem_col).value =f'Total coded OMRs'
    coding_rem_sheet.cell(row=4, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=4, column=coding_rem_col).font = Font(size = 12)
    #coding_rem_sheet.cell(row=4, column=1).border = thin_border
    coding_rem_sheet.cell(row=4, column=coding_rem_col).alignment = center_alignment_obj


    coding_rem_sheet.cell(row=7, column=coding_rem_col).value =f'=SUM(F7:{get_column_letter(coding_rem_col-1)}7)'
    coding_rem_sheet.cell(row=7, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=7, column=coding_rem_col).font = Font(size = 12)
    
    coding_rem_sheet.cell(row=8, column=coding_rem_col).value =f'Total coding Remuneration'
    coding_rem_sheet.cell(row=7, column=coding_rem_col).font = Font(bold=True)
    coding_rem_sheet.cell(row=7, column=coding_rem_col).font = Font(size = 12)
    coding_rem_sheet.cell(row=7, column=coding_rem_col).alignment = wrap_text__alignment_obj

    for i in range(6,inv_rem_col):
        coding_rem_sheet.cell(row=31, column= i).value =f'=SUM({get_column_letter(i)}9:{get_column_letter(i)}30)'

    coding_rem_sheet.merge_cells(start_row=31, start_column=1, end_row=31, end_column=5)

    for i in range(2,24):
        formula2 = ''
        for j in range(6,coding_rem_col):
            if j == coding_rem_col-1:

                formula2 = formula2 + f'{get_column_letter(j)}{coding_rem_row+3+i}*IF(ISERROR({get_column_letter(5)}{coding_rem_row+3+i}/{get_column_letter(j)}31),0,{get_column_letter(5)}{coding_rem_row+3+i}/{get_column_letter(j)}31)*{get_column_letter(j)}7' 

            else:
                formula2 = formula2 + f'{get_column_letter(j)}{coding_rem_row+3+i}*IF(ISERROR({get_column_letter(5)}{coding_rem_row+3+i}/{get_column_letter(j)}31),0,{get_column_letter(5)}{coding_rem_row+3+i}/{get_column_letter(j)}31)*{get_column_letter(j)}7 + ' 
        #print(formula2+'\n')
        coding_rem_sheet.cell(row=coding_rem_row+3+i, column=coding_rem_col).value =f'=ROUNDUP({formula2},0)'

        coding_rem_sheet.cell(row=31, column=coding_rem_col).value =f'=sum({get_column_letter(coding_rem_col)}9:{get_column_letter(coding_rem_col)}30)'

        coding_rem_sheet.merge_cells(start_row=32, start_column=1, end_row=32, end_column=coding_rem_col-1)

        coding_rem_sheet.cell(row=32, column=1).value =f'Total:'
        coding_rem_sheet.cell(row=32, column=1).alignment = right_alignment_obj
        coding_rem_sheet.cell(row=32, column=coding_rem_col).value =f'=ROUNDUP({get_column_letter(coding_rem_col)}7*0.25,0)'


        coding_rem_sheet.merge_cells(start_row=33, start_column=1, end_row=36, end_column=coding_rem_col)
        coding_rem_sheet.cell(row=33, column=1).value =f'PRINCIPAL'  
        coding_rem_sheet.cell(row=33, column=1).font = Font(bold=True)
        #coding_rem_sheet.cell(row=30, column=1).border = thin_border
        coding_rem_sheet.cell(row=33, column=1).alignment = right_alignment_obj









    for col in inv_sheet.columns:
        for cell in col:
            # openpyxl styles aren't mutable,
            # so you have to create a copy of the style, modify the copy, then set it back
            
            cell.border = thin_border

    for col in inv_rem_sheet.columns:
        for cell in col:
            # openpyxl styles aren't mutable,
            # so you have to create a copy of the style, modify the copy, then set it back
            
            cell.border = thin_border

    for col in coding_rem_sheet.columns:
        for cell in col:
            # openpyxl styles aren't mutable,
            # so you have to create a copy of the style, modify the copy, then set it back
            
            cell.border = thin_border


    inv_sheet.column_dimensions['B'].width = 30
    inv_sheet.column_dimensions['C'].width = 20
    inv_sheet.column_dimensions['D'].width = 10
    invisilation_list.save(op_file_path+'\\Invisilation_Duties_Proforma.xlsx')

    inv_rem_sheet.column_dimensions['B'].width = 30
    inv_rem_sheet.column_dimensions['C'].width = 20
    inv_rem_sheet.column_dimensions['D'].width = 10
    inv_rem_sheet.column_dimensions['E'].width = 15
    invisilation_rem.save(op_file_path+'\\Invisilation_Remunerations.xlsx')

    coding_rem_sheet.column_dimensions['B'].width = 30
    coding_rem_sheet.column_dimensions['C'].width = 20
    coding_rem_sheet.column_dimensions['D'].width = 10
    coding_rem_sheet.column_dimensions['E'].width = 15
    coding_rem.save(op_file_path+'\\Coding(Tearing)_Remunerations.xlsx')

    print("done!")

def qp_opening_fill(collegeName,name_of_Exam):
    qp_opening = Document()
    section = qp_opening.sections[0]
    #section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    sections = qp_opening.sections
    for section in sections:
                margin = 2
                section.top_margin = Cm(margin)
                section.bottom_margin = Cm(margin)
                section.left_margin = Cm(margin)
                section.right_margin = Cm(margin)
                
    #print(date_list)
    print("Creating Summary for Opening Certficate .......")
    qp_opening.add_heading(f'{collegeName}\n{name_of_Exam}\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p=qp_opening.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(f'Opening Certificates')

    runner.bold = True
    runner.font.size = Pt(90)

    qp_opening.add_page_break()
    for date in date_list:
        for session in session_list:
            if session == 'AN':
                
                tim = ' 2:30 PM to 4:30 PM '
                
            elif session == 'FN':
                tim = ' 9:30 AM to 11:30 AM '
            elif session == 'BN':
                tim = ' 12:00 PM to 02:00 AM '
            else:
                tim = ''
            
            print(f"Filling Summary for Opening Certificate for DATE: {date} and Session: {session}............",end='')
            print('Done!')
            
            
            #print(f'        Date: {date}       Session : {session}')
            #document.add_heading(f'        Date: {date}       Session : {session}')
            #qp_opening.add_heading(f'Opening Certificate Summary for Date: {strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}       Session : {session}')
            k=qp_opening.add_paragraph()
            k.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            runner = p.add_run(f'Opening Certificate\n')
            runner.font.underline = WD_UNDERLINE.SINGLE

            runner.bold = True
            runner.font.size = Pt(14)

            # qp_opening.add_heading(f'Opening Certificate\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # runner.font.underline = WD_UNDERLINE.SINGLE

                        
            
            table = qp_opening.add_table(rows=1, cols= 4)

            set_column_width(table.columns[0], Cm(1))
            set_column_width(table.columns[1], Cm(6.5))
            set_column_width(table.columns[2], Cm(.5))
            set_column_width(table.columns[3], Cm(8.5))

            
           
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '1.'
            hdr_cells[1].text = 'Examination Centre'
            hdr_cells[2].text = ':'
            hdr_cells[3].text = f'{collegeName}'

            hdr_cells = table.add_row().cells
            hdr_cells[0].text = '2.'
            hdr_cells[1].text = 'Month & Year of Examination'
            hdr_cells[2].text = ':'
            hdr_cells[3].text = f' {strftime("%B",strptime(date,"%Y-%m-%d"))}, {strftime("%Y",strptime(date,"%Y-%m-%d"))}'.upper()

            hdr_cells = table.add_row().cells
            hdr_cells[0].text = '3.'
            hdr_cells[1].text = 'Day & Date of Examination'
            hdr_cells[2].text = ':'
            hdr_cells[3].text = f'{strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}, {strftime("%A",strptime(date,"%Y-%m-%d"))}'

            hdr_cells = table.add_row().cells
            hdr_cells[0].text = '4.  '
            hdr_cells[1].text = 'Time of Examination'
            hdr_cells[2].text = ':'                                 
            hdr_cells[3].text = f'{tim}'

            
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            #table.style = 'Light Shading'
            #table.style = 'Table Grid'
            #document.save(file_path+'/op.docx')
            sc = ''
            sn = ''
            i=0
            for scheme in scheme_list:

                for dept in branch_list:
                    command = f'''SELECT year,sub_code,sub_name,sub_code_no FROM {time_table} WHERE date LIKE '{date}' 
                                                                                                    AND session LIKE '{session}'
                                                                                                    AND branch LIKE '{dept}'
                                                                                                    AND scheme LIKE '{scheme}';'''
                    #print(command)
                    curs.execute(command)
                    a= curs.fetchall()
                    if len(a) > 0:
                        for (year,sub_code,sub_name,sub_code_no) in a:
                            if i == 0:
                                sc = f'{scheme}_{sub_code}'
                                sn = f"{sub_name}"
                            else:
                                sc += f', {scheme}_{sub_code}'
                                sn += f", {sub_name}"
                            i+=1
                    

            hdr_cells = table.add_row().cells
            hdr_cells[0].text = '5.  '
            hdr_cells[1].text = 'Subject Code Nos'
            hdr_cells[2].text = ':'
            if len(sc)>0:
                hdr_cells[3].text = f'{sc}'
            if len(sc)==0:
                hdr_cells[3].text = f'--NIL--'


            hdr_cells = table.add_row().cells
            hdr_cells[0].text = '6.  '
            hdr_cells[1].text = 'Name of the Subjects'
            hdr_cells[2].text = ':'
            if len(sn)>0:
                hdr_cells[3].text = f'{sn}'
            if len(sn)==0:
                hdr_cells[3].text = f'--NIL--'


            

            doc2_cells = table.add_row().cells
                                        #print(slno)
                        
            merged_cell = doc2_cells[0].merge(doc2_cells[1]).merge(doc2_cells[2]).merge(doc2_cells[3])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            merged_cell.text = '\tWe the undersigned hereby certify that the sealed envelope containing question paper in the subject mentioned above have been closely examined by us and found to be in proper condition with the seals intact. The envelope has been opened in the Examination hall in our presence on the due date and time. We further certify that the copies of question papers in the envelope has been carefully counted and number of copies was found to be correct as indicated on the envelope.'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.JUSTIFY
            merged_cell.paragraphs[0].runs[0].font.bold = False

            doc2_cells = table.add_row().cells
            merged_cell1 = doc2_cells[0].merge(doc2_cells[1])
            merged_cell1.alignment=WD_ALIGN_PARAGRAPH.LEFT
            merged_cell1.text = f'Date: {strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}'
            merged_cell2 = doc2_cells[2].merge(doc2_cells[3])
            merged_cell2.text = f''' \nChief Superintendent \nName & Designation \n(\t\t\t) \n\nJoint Chief Superintendent \nName & Designation \n(\t\t\t) \n\nInvisilator-1 \nName & Designation \n(\t\t\t) \n\nInvisilator-2 \nName & Designation \n(\t\t\t) '''
            merged_cell2.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER                                   


                                   
            qp_opening.add_page_break()
            #sheet.column_dimensions.hidden = False

            

    
    print("Creating EDEP summary.........",end='')
    os.makedirs(op_file_path, exist_ok=True)
    qp_opening.save(op_file_path+'\\OpeningCerificates.docx')
    print("done!")



def EDEP_fill(collegeName,name_of_Exam):
    EDEP_document = Document()
    section = EDEP_document.sections[0]
    #section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    sections = EDEP_document.sections
    for section in sections:
                margin = 2
                section.top_margin = Cm(margin)
                section.bottom_margin = Cm(margin)
                section.left_margin = Cm(margin)
                section.right_margin = Cm(margin)
                
    #print(date_list)
    print("Creating Summary for EDEP .......")
    EDEP_document.add_heading(f'{collegeName}\n{name_of_Exam}\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p=EDEP_document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(f'ACCOUNT OF EDEP')
    runner.bold = True
    runner.font.size = Pt(104)

    EDEP_document.add_page_break()

    for date in date_list:
        for session in session_list:
            print(f"Filling Summary for EDEP for DATE: {date} and Session: {session}............",end='')
            print('Done!')
            grand_total = 0
            slno=0
            total_qps = 0
            
            #print(f'        Date: {date}       Session : {session}')
            #document.add_heading(f'        Date: {date}       Session : {session}')
            EDEP_document.add_heading(f'{collegeName}\n{name_of_Exam}').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            EDEP_document.add_heading(f'EDEP Summary for Date: {strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}       Session : {session}').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            table1 = EDEP_document.add_table(rows=1, cols=6)
            table1.autofit = False
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells

            hdr_cells[0].text = f'Password Received Time:'
            hdr_cells[1].text = f''
            hdr_cells[2].text = f'QP Printing start Time:'
            hdr_cells[3].text = f''' '''
            hdr_cells[4].text = f'''QP Printing End Time:'''
            hdr_cells[5].text = f''' '''

            set_column_width(table1.columns[0], Cm(4))
            set_column_width(table1.columns[1], Cm(1))
            set_column_width(table1.columns[2], Cm(3))
            set_column_width(table1.columns[3], Cm(1))
            set_column_width(table1.columns[4], Cm(3))
            set_column_width(table1.columns[5], Cm(1))
            
            
            topRow = ['S.No','Scheme','Year','Subject Code','Subject Name','Strength','No. Q.P. Printed']
            table = EDEP_document.add_table(rows=1, cols=len(topRow))

            #topRow = ['S.No','Scheme','Year','Subject Code','Subject Name','Strength','No. Q.P. to be Printed']
            hdr_cells = table.rows[0].cells
            for i in range(0,len(hdr_cells)):
                                hdr_cells[i].text = topRow[i]
                                run = hdr_cells[i].paragraphs[0].runs[0]
                                run.font.bold = True
            table.autofit = False

            #table.style = 'Light Shading'
            table.style = 'Table Grid'
            #document.save(file_path+'/op.docx')
            for scheme in scheme_list:

                for dept in branch_list:
                    command = f'''SELECT year,sub_code,sub_name,sub_code_no FROM {time_table} WHERE date LIKE '{date}' 
                                                                                                    AND session LIKE '{session}'
                                                                                                    AND branch LIKE '{dept}'
                                                                                                    AND scheme LIKE '{scheme}';'''
                    #print(command)
                    curs.execute(command)
                    a= curs.fetchall()
                    
                    if len(a) > 0:
                        for (year,sub_code,sub_name,sub_code_no) in a:
                            slno+=1
                            #print(year,sub_code,sub_name,sub_code_no)
                             
                            command = f'''SELECT PinNo FROM {nr_table} WHERE scheme LIKE '{scheme}'
                                                                        AND branch LIKE '{dept}'
                                                                        AND year LIKE '{year}'
                                                                        AND (SUB01 LIKE '%{sub_code_no}%'
                                                                        OR SUB02 LIKE '%{sub_code_no}%'
                                                                        OR SUB03 LIKE '%{sub_code_no}%'
                                                                        OR SUB04 LIKE '%{sub_code_no}%'
                                                                        OR SUB05 LIKE '%{sub_code_no}%'
                                                                        OR SUB06 LIKE '%{sub_code_no}%'
                                                                        OR SUB07 LIKE '%{sub_code_no}%'
                                                                        OR SUB08 LIKE '%{sub_code_no}%'
                                                                        OR SUB09 LIKE '%{sub_code_no}%'
                                                                        OR SUB10 LIKE '%{sub_code_no}%'
                                                                        OR SUB11 LIKE '%{sub_code_no}%'
                                                                        OR SUB12 LIKE '%{sub_code_no}%')
                                                                        ORDER BY PinNo ASC;'''
                            
                                
                            curs.execute(command)
                            pins = curs.fetchall()
                            pins = elDupes(pins)
                            #print(pins)
                            pins.sort()

                            row_cells = table.add_row().cells
                            row_cells[0].text = f'{slno}'
                            row_cells[1].text = f'{scheme}'
                            row_cells[2].text = f'{year}'
                            row_cells[3].text = f'''{sub_code}'''
                            row_cells[4].text = f'''{sub_name}'''                           
                            
                            if len(pins) >0:
                                                            qps =len(pins)+2
                                                            row_cells[5].text = noneType(len(pins))
                                                            row_cells[6].text = f'''{noneType(qps)}'''
                            if len(pins)== 0:
                                                            qps = 0
                                                            row_cells[5].text = f'--NIL--'
                                                            row_cells[6].text = f'''--NIL--'''
                            

                            grand_total += len(pins)
                            total_qps += qps

            row_cells = table.add_row().cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.text = 'Total:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True

            row_cells[5].text = f'{str(grand_total).zfill(2)}'
            row_cells[5].paragraphs[0].runs[0].font.bold = True
            row_cells[6].text = f'{str(total_qps).zfill(2)}'
            row_cells[6].paragraphs[0].runs[0].font.bold = True


            tableRow = table.add_row()
            tableRow.height_rule = WD_ROW_HEIGHT.EXACTLY
            tableRow.height = Pt(30)
            row_cells = tableRow.cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell1=row_cells[5].merge(row_cells[6])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.text = 'Signature of EDEP Incharge:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True

            tableRow = table.add_row()
            tableRow.height_rule = WD_ROW_HEIGHT.EXACTLY
            tableRow.height = Pt(30)

            row_cells = tableRow.cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell1=row_cells[5].merge(row_cells[6])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.text = 'Signature of JCS:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True

            tableRow = table.add_row()
            tableRow.height_rule = WD_ROW_HEIGHT.EXACTLY
            tableRow.height = Pt(30)

            row_cells = tableRow.cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell1=row_cells[5].merge(row_cells[6])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.text = 'Signature of CS:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True
            
            
            #row_cells[7].text = ' '

                            

            
            
            set_column_width(table.columns[0], Cm(1.6))
            set_column_width(table.columns[1], Cm(2.5))
            set_column_width(table.columns[2], Cm(1.5))
            set_column_width(table.columns[3], Cm(2.5))
            set_column_width(table.columns[4], Cm(5))
            set_column_width(table.columns[5], Cm(2.5))
            set_column_width(table.columns[6], Cm(2))

            
            
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(12)
            #EDEP_document.add_heading(f'                          Total Strength : {grand_total}')
            #EDEP_document.add_heading(f'      Total No .of question papers printed : {total_qps}')
            #paragraph = EDEP_document.add_paragraph()
            #paragraph.text = '''           Developed by Keerthi Chandra C, Lecturer in ECE, GMRPW Karimnagar
            #                           reach me at keerthichand.c@gmail.com'''



            EDEP_document.add_page_break()
            #sheet.column_dimensions.hidden = False

            

    
    print("Creating EDEP summary.........",end='')
    os.makedirs(op_file_path, exist_ok=True)
    EDEP_document.save(op_file_path+'\\EDEP_Summary.docx')
    print("done!")


def QP_account_fill(collegeName,name_of_Exam):
    QP_account = Document()
    section = QP_account.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    sections = QP_account.sections
    for section in sections:
                margin = 2
                section.top_margin = Cm(margin)
                section.bottom_margin = Cm(margin)
                section.left_margin = Cm(margin)
                section.right_margin = Cm(margin)
                
    #print(date_list)
    print("Creating Question Paper Account .......")

    QP_account.add_heading(f'{collegeName}\n{name_of_Exam}\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p=QP_account.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = p.add_run(f'Question Paper Balance Account')
    runner.bold = True
    runner.font.size = Pt(95)
    QP_account.add_page_break()


    for date in date_list:
        for session in session_list:
            print(f"Filling QP account for DATE: {date} and Session: {session}............",end='')
            print('Done!')
            grand_total = 0
            slno=0
            total_qps = 0
            
            #print(f'        Date: {date}       Session : {session}')
            #document.add_heading(f'        Date: {date}       Session : {session}')
            QP_account.add_heading(f'{collegeName}\n{name_of_Exam}').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            QP_account.add_heading(f'Question Paper Account for   Date: {strftime("%d-%B-%Y",strptime(date,"%Y-%m-%d"))}       Session : {session}')


            table = QP_account.add_table(rows=1, cols=9)
            topRow = ['S.No','Scheme','Year','Subject Code','Subject Name','Strength','No. Q.P. Received','No. Q.P. Used','No. Q.P. Balance']
            hdr_cells = table.rows[0].cells
            for i in range(0,len(hdr_cells)):
                                hdr_cells[i].text = topRow[i]
                                run = hdr_cells[i].paragraphs[0].runs[0]
                                run.font.bold = True
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            #table.style = 'Light Shading'
            table.style = 'Table Grid'
            #document.save(file_path+'/op.docx')
            for scheme in scheme_list:

                for dept in branch_list:
                    command = f'''SELECT year,sub_code,sub_name,sub_code_no FROM {time_table} WHERE date LIKE '{date}' 
                                                                                                    AND session LIKE '{session}'
                                                                                                    AND branch LIKE '{dept}'
                                                                                                    AND scheme LIKE '{scheme}';'''
                    #print(command)
                    curs.execute(command)
                    a= curs.fetchall()
                    
                    if len(a) > 0:
                        for (year,sub_code,sub_name,sub_code_no) in a:
                            slno+=1
                            #print(year,sub_code,sub_name,sub_code_no)
                             
                            command = f'''SELECT PinNo FROM {nr_table} WHERE scheme LIKE '{scheme}'
                                                                        AND branch LIKE '{dept}'
                                                                        AND year LIKE '{year}'
                                                                        AND (SUB01 LIKE '%{sub_code_no}%'
                                                                        OR SUB02 LIKE '%{sub_code_no}%'
                                                                        OR SUB03 LIKE '%{sub_code_no}%'
                                                                        OR SUB04 LIKE '%{sub_code_no}%'
                                                                        OR SUB05 LIKE '%{sub_code_no}%'
                                                                        OR SUB06 LIKE '%{sub_code_no}%'
                                                                        OR SUB07 LIKE '%{sub_code_no}%'
                                                                        OR SUB08 LIKE '%{sub_code_no}%'
                                                                        OR SUB09 LIKE '%{sub_code_no}%'
                                                                        OR SUB10 LIKE '%{sub_code_no}%'
                                                                        OR SUB11 LIKE '%{sub_code_no}%'
                                                                        OR SUB12 LIKE '%{sub_code_no}%')
                                                                        ORDER BY PinNo ASC;'''
                            
                                
                            curs.execute(command)
                            pins = curs.fetchall()
                            pins = elDupes(pins)
                            #print(pins)
                            pins.sort()

                            row_cells = table.add_row().cells
                            row_cells[0].text = f'{slno}'
                            row_cells[1].text = f'{scheme}'
                            row_cells[2].text = f'{year}'
                            row_cells[3].text = f'''{sub_code}'''
                            row_cells[4].text = f'''{sub_name}'''                           
                            row_cells[5].text = noneType(len(pins))
                            if len(pins) >0:
                                                            qps =len(pins)+2
                                                            row_cells[5].text = noneType(len(pins))
                                                            row_cells[6].text = f'''{noneType(qps)}'''
                            if len(pins)== 0:
                                                            qps = 0
                                                            row_cells[5].text = f'--NIL--'
                                                            row_cells[6].text = f'''--NIL--'''
                            

                            grand_total += len(pins)
                            total_qps += qps

                            

            
            row_cells = table.add_row().cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.text = 'Total:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True
            
            row_cells[5].text = f'{str(grand_total).zfill(2)}'
            row_cells[5].paragraphs[0].runs[0].font.bold = True
            row_cells[6].text = f'{str(total_qps).zfill(2)}'
            row_cells[6].paragraphs[0].runs[0].font.bold = True
            row_cells[7].text = ' '

            tableRow = table.add_row()
            tableRow.height_rule = WD_ROW_HEIGHT.EXACTLY
            tableRow.height = Pt(30)
            row_cells = tableRow.cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell1 = row_cells[5].merge(row_cells[6]).merge(row_cells[7]).merge(row_cells[8])
            merged_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.text = 'Signature of AS-1:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True


            tableRow = table.add_row()
            tableRow.height_rule = WD_ROW_HEIGHT.EXACTLY
            tableRow.height = Pt(30)
            row_cells = tableRow.cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell1 = row_cells[5].merge(row_cells[6]).merge(row_cells[7]).merge(row_cells[8])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.text = 'Signature of AS-2:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True




            tableRow = table.add_row()
            tableRow.height_rule = WD_ROW_HEIGHT.EXACTLY
            tableRow.height = Pt(30)
            row_cells = tableRow.cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell1 = row_cells[5].merge(row_cells[6]).merge(row_cells[7]).merge(row_cells[8])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.text = 'Signature of JCS:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True

            tableRow = table.add_row()
            tableRow.height_rule = WD_ROW_HEIGHT.EXACTLY
            tableRow.height = Pt(30)
            row_cells = tableRow.cells
            merged_cell = row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3]).merge(row_cells[4])
            merged_cell1 = row_cells[5].merge(row_cells[6]).merge(row_cells[7]).merge(row_cells[8])
            merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            merged_cell.text = 'Signature of CS:'
            #print(merged_cell.paragraphs[0])
            merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
            merged_cell.paragraphs[0].runs[0].font.bold = True

            set_column_width(table.columns[0], Cm(1.1))
            set_column_width(table.columns[1], Cm(1.8))
            set_column_width(table.columns[2], Cm(1.5))
            set_column_width(table.columns[3], Cm(1.8))
            set_column_width(table.columns[4], Cm(5))
            set_column_width(table.columns[5], Cm(1.8))
            set_column_width(table.columns[6], Cm(2))
            set_column_width(table.columns[7], Cm(2))
            set_column_width(table.columns[8], Cm(2))
            #set_column_width(table.columns[9], Cm(1))
            
            for row in table.rows:
                for cell in row.cells:
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        for run in paragraph.runs:
                            font = run.font
                            font.size= Pt(10)
            QP_account.add_heading(f'''

''')
            #QP_account.add_heading(f' AS-1 \t\t\t\t AS-2 \t\t\t\t JCS \t\t\t\t CS')
            
            paragraph = QP_account.add_paragraph()
            #paragraph.text = '''           Developed by Keerthi Chandra C, Lecturer in ECE, GMRPW Karimnagar
            #                           reach me at keerthichand.c@gmail.com'''



            QP_account.add_page_break()
            #sheet.column_dimensions.hidden = False

            

    
    print("Creating QP_Account.........",end='')
    os.makedirs(op_file_path, exist_ok=True)
    QP_account.save(op_file_path+'\\QP_account.docx')
    print("done!")
    


if __name__ == '__main__':
    
    coderintro()
    instructions_info()
    #create_col_info()
    while(True):
        print("Getting College details......")                
        collegeCode = get_from_DB("collegeCode",COLLEGE_INFO_TABLE)
        
        if len(collegeCode) == 0:
            
            print("College Details Not Found...........!")
            update_collInfo()
        collegeCode = get_from_DB("collegeCode",COLLEGE_INFO_TABLE)
        collegeName = get_from_DB("collegeName",COLLEGE_INFO_TABLE)
        if len(collegeName) == 0:
            
            print("College Details Not Entered...........!")
            update_collInfo()
        else:
            collegeCode = get_from_DB("collegeCode",COLLEGE_INFO_TABLE)
            collegeName = get_from_DB("collegeName",COLLEGE_INFO_TABLE)
            break
    collegeCode=collegeCode[0]
    collegeName = collegeName[0]
    print(collegeCode)
    print(collegeName)


    
    
    
    
    
    options = [' 1. Create Pre Exam Formats',' 2. Generate DE-4 Attendance,Packing slips and Balance Account',' 3. Update Time Tables',' 4. Update College Information']
    
    
    while(True):
        print()
        print("**********************************************************************************************************")
        print()
        for option in options:
            print(option)
        print()
        print("**********************************************************************************************************")
        print()
        sel= int(input(f" Enter your option: "))
        if sel == 2:
                        #PackingSlip = Document(file_path+'/srcs/Template_PackingSlip.docx')
                        PackingSlip = Document()

                        section = PackingSlip.sections[0]
                        section.page_height = Cm(29.7)
                        section.page_width = Cm(21)
                        sections = PackingSlip.sections

                        for section in sections:
                            margin =0.5
                            margin1 = 2
                            section.top_margin = Cm(margin)
                            section.bottom_margin = Cm(margin)
                            section.left_margin = Cm(margin1)
                            section.right_margin = Cm(margin1)



                        PackingSlip.add_heading('PROFORMA – III\n').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p=PackingSlip.add_paragraph()
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        runner = p.add_run(f'PACKING SLIP: ')
                        runner.font.underline = WD_UNDERLINE.SINGLE
                        runner.bold = True
                        runner.font.size = Pt(14)


                        table1 = PackingSlip.add_table(rows=3,cols=3)
                        table1.autofit = False
                        table1.alignment = WD_TABLE_ALIGNMENT.CENTER
                        set_column_width(table1.columns[0], Cm(9))
                        set_column_width(table1.columns[1], Cm(.5))
                        set_column_width(table1.columns[2], Cm(8))

                        PackingSlip.add_paragraph('''\nDetails of Courses in Insured Parcel:''')

                        do_cells = table1.rows[0].cells
                        do_cells[0].text = 'Name of the Institution & Code No. '
                        do_cells[1].text = ':'
                        do_cells[2].text = f' {collegeCode}, {collegeName}'

                        table2 = PackingSlip.add_table(rows=1, cols=8)
                        table2.autofit = False
                        table2.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table2.style = 'Light Shading'
                        #table2.style = 'TableGrid'
                        set_column_width(table2.columns[0], Cm(1.5))
                        set_column_width(table2.columns[1], Cm(2.8))
                        set_column_width(table2.columns[2], Cm(1.8))
                        set_column_width(table2.columns[3], Cm(2))
                        set_column_width(table2.columns[4], Cm(2))
                        set_column_width(table2.columns[5], Cm(2))
                        set_column_width(table2.columns[6], Cm(2))
                        set_column_width(table2.columns[7], Cm(2))

                        doc2_cells = table2.rows[0].cells

                        doc2_cells[0].text = '{0}'.format('Sl. No')
                        doc2_cells[1].text = '{0}'.format('Department')
                        doc2_cells[2].text = '{0}'.format("scheme")
                        doc2_cells[3].text = '{0}'.format('Subject Code')
                        doc2_cells[4].text = '{0}'.format('Year/Semester')
                        doc2_cells[5].text = '{0}'.format('Number of Scripts')
                        doc2_cells[6].text = '{0}'.format('Number of Packets')
                        doc2_cells[7].text = '{0}'.format('Remarks')

                        print("Getting Exam Dates......",end='')                
                        date_list = get_from_DB("date",time_table)
                        print('done!')
                        session_list = get_from_DB("session",time_table)
                        print("Getting Branch List......",end='')
                        branch_list = get_from_DB("branch",nr_table)
                        print('done!')
                        scheme_list = get_from_DB("scheme",nr_table)

    #conn.close()
                        date_list = elDupes(date_list)
                        date_list.sort()
                        session_list = elDupes(session_list)
                        session_list.sort(reverse = True)
                        scheme_list = elDupes(scheme_list)
                        scheme_list.sort()
                        branch_list = elDupes(branch_list)
                        branch_list.sort()
                        date,session= dateAndSession()
                        #date,session= '2018-10-30','AN'
                        sl=0
                        number_0f_packets = 0
                        number_0f_scripts = 0
                        for scheme in scheme_list:

                                        for dept in branch_list:
                                             sl,number_0f_scripts,number_0f_packets,pathe =   deFill(sl,number_0f_scripts,number_0f_packets,date,session,scheme,dept,collegeName,collegeCode)
                        #PackingSlip.add_heading(f'      Total No. of Packtes : {number_0f_packets}({numtoWord(number_0f_packets)})')
                        doc2_cells = table2.add_row().cells
                                        #print(slno)
                        
                        merged_cell = doc2_cells[0].merge(doc2_cells[1]).merge(doc2_cells[2]).merge(doc2_cells[3]).merge(doc2_cells[4])
                        merged_cell.alignment=WD_ALIGN_PARAGRAPH.RIGHT
                        merged_cell.text = 'Total'
                        #print(merged_cell.paragraphs[0])
                        merged_cell.paragraphs[0].alignment  = WD_ALIGN_PARAGRAPH.RIGHT
                        merged_cell.paragraphs[0].runs[0].font.bold = True
                        
                        doc2_cells[5].text = f'{str(number_0f_scripts).zfill(2)}'
                        doc2_cells[5].paragraphs[0].runs[0].font.bold = True
                        doc2_cells[6].text = f'{str(number_0f_packets).zfill(2)}'
                        doc2_cells[6].paragraphs[0].runs[0].font.bold = True
                        doc2_cells[7].text = ' '
                        
                        paragraph = PackingSlip.add_paragraph()
                        paragraph.text = f'''\n                                                   Total No. of Packets: {str(number_0f_packets).zfill(2)} ({numtoWord(number_0f_packets )}ONLY)\n\n\n\n\n                                                                                                          SIGNATURE OF CHIEF SUPERINTENDENT'''

                                            
                        PackingSlip.save(pathe+'/Packing Slip.docx')
                        print("all files are saved in OUTPUTS folder")
                        break
                                                

        
        if sel == 1:
                        print("Getting Exam Dates......",end='')                
                        date_list = get_from_DB("date",time_table)
                        print('done!')
                        session_list = get_from_DB("session",time_table)
                        print("Getting Branch List......",end='')
                        print("Reading NRs.......")
                        create_nr()
                        nr_DB()
                        print("Done!")
                        branch_list = get_from_DB("branch",nr_table)
                        print('done!')
                        scheme_list = get_from_DB("scheme",nr_table)

    #conn.close()
                        date_list = elDupes(date_list)
                        date_list.sort()
                        
                                
                        months_list = []
                        years_list = []
                        for date in date_list:
                            months_list.append(strftime("%m",strptime(date,"%Y-%m-%d")).upper())
                            years_list.append(strftime("%Y",strptime(date,"%Y-%m-%d")).upper())

                        months_list = elDupes(months_list)
                        years_list = elDupes(years_list)
                        months_list.sort()
                        years_list.sort()

                        #print(months_list,years_list)

                        name_of_Exam = 'SBTET Examinations'

                        k = 0
                        for month in months_list:
                            if k == 0:
                                name_of_Exam = f'{name_of_Exam} {strftime("%b",strptime(month,"%m")).upper()}'
                                k+=1
                            else:
                                name_of_Exam = f'{name_of_Exam}/{strftime("%b",strptime(month,"%m")).upper()}'
                        k = 0
                        for year in years_list:
                            if k == 0:
                                name_of_Exam = f'{name_of_Exam}-{year}'
                                k+=1
                            else:
                                name_of_Exam = f'{name_of_Exam}/{month}'

                        #print(name_of_Exam)

                        session_list = elDupes(session_list)
                        session_list.sort(reverse = True)
                        scheme_list = elDupes(scheme_list)
                        scheme_list.sort()
                        branch_list = elDupes(branch_list)
                        branch_list.sort()
                        #create_nr()
                        #nr_DB()
                        summary_fill(collegeCode,collegeName,name_of_Exam)
                        EDEP_fill(collegeName,name_of_Exam)
                        QP_account_fill(collegeName,name_of_Exam)
                        qp_opening_fill(collegeName,name_of_Exam)
                        print("all files are saved in OUTPUTS folder")
                        break
        if sel == 3:
                        create_TT()
                        update_TT()
                        break

        if sel == 4:
                        update_collInfo()
                        break
                        

        else:
            print("Wrong Option Entered!!\nEnter Again......")
            


    conn.close()
    #doc2.save(op_file_path+'/Packing Slip.docx')
    input("press Enter to exit.....")
